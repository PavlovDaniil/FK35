from docxtpl import DocxTemplate
from docx import Document
from docx2pdf import convert
import re
import uuid
import os

PLACEHOLDER_PATTERN = r"{{(.*?)}}"
PATH_TEMPLATE = "templates/"
OUTPUT_DIR = "out/"


def generate_pdf_from_template(file_name: str, data: dict):
    # Загружаем шаблон
    doc = DocxTemplate(f"{PATH_TEMPLATE}{file_name}")

    # Подставляем данные
    doc.render(data)

    # Имя временного docx
    temp_docx = os.path.join(OUTPUT_DIR, f"{uuid.uuid4()}.docx")
    temp_pdf = temp_docx.replace(".docx", ".pdf")

    # Сохраняем docx
    doc.save(temp_docx)

    # Конвертация в PDF
    convert(temp_docx, temp_pdf)

    # Можно удалить промежуточный docx
    os.remove(temp_docx)

    return temp_pdf

def search(template_name: str) -> set:
    text = read_file(template_name)

    # ищем плейсхолдеры
    placeholders = set(re.findall(PLACEHOLDER_PATTERN, text))
    return placeholders

def get_all_files() -> list:
    return os.listdir(PATH_TEMPLATE)

def read_file(file_name: str) -> str:
    doc = Document(f"{PATH_TEMPLATE}{file_name}")

    text = ""

    # собираем текст из всех параграфов и таблиц
    for para in doc.paragraphs:
        text += para.text + "\n"

    # таблицы тоже обрабатываем
    for table in doc.tables:
        for row in table.rows:
            for cell in row.cells:
                text += cell.text + "\n"

    return text